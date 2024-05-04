import fitz
from io import BytesIO
import docx
import re
from custom.model import model
import zipfile
import asyncio
import json

def read_docx(file_bytes):
    doc = docx.Document(BytesIO(file_bytes))
    all_para = []
    for para in doc.paragraphs:
        all_para.append(para.text)
    # all_para = "\n".join(all_para)
    return all_para

def read_pdf(file_bytes):
    pdf_file = fitz.open('pdf',file_bytes)
    all_para = []
    for i, page in enumerate(pdf_file):
        # data.append(page.get_text())
        page_text = page.get_text() 
        if page_text.strip() == '':
            print(f"nat able to extract text from page {i+1}")
        all_para.extend(page_text.split("\n"))
        
    return all_para


def write_docx(result_text):
    document = docx.Document()

    document.add_paragraph(result_text)
    target_stream = BytesIO()
    document.save(target_stream)
    return target_stream

def zip_files(file_tuples):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        for file_name, bytes_io in file_tuples:
            zip_file.writestr(file_name, bytes_io.getvalue())
    zip_buffer.seek(0)
    return zip_buffer


async def make_qa_from_para(para):
    # print(para)
    if para.strip() == '':
        return None

    
    response = await model.generate_content(f"""
        Given a Paragraph make one Question and answer based on that paragraph.
        The question should be objective with only one answer.
        The answer should be one to three words.
        
        input para: {para}    

        output format:
        ```
        {{
            "question": str,
            "answer": str
        }}
        ```
        
        Start and end the output with ```
    """)
    # print(res)
    try:
        json_str = re.search(r'```\n(.*)\n```', response, re.DOTALL)[1]
        data = json.loads(json_str)
    except Exception as e:
        print("unable to decode response\n\n",response, '\n\n')
        raise e

    
    data['para'] = para
    return data


async def make_qa_from_all_para(all_para):
    tasks = [make_qa_from_para(para) for para in all_para]
    results = await asyncio.gather(*tasks)
    return [r for r in results if r]

